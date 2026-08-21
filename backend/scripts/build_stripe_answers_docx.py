"""
Generate a clean, send-ready DOCX response to CypherPro's Stripe/NMI
migration questions. Every answer has been reviewed for accuracy
against the codebase and rewritten where the original was ambiguous,
contradictory, or missing.

Output: /app/frontend/public/downloads/CypherPro_Stripe_Answers.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/app/frontend/public/downloads/CypherPro_Stripe_Answers.docx")

# ---- Brand palette
NAVY = RGBColor(0x0F, 0x17, 0x2A)
CYAN = RGBColor(0x08, 0x91, 0xB2)
SLATE = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# ---- Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = CYAN
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def q(number, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"Q{number}. ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = SLATE
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def a(text):
    p = doc.add_paragraph()
    r = p.add_run("Answer. ")
    r.bold = True
    r.font.color.rgb = CYAN
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

def spacer():
    doc.add_paragraph("")


# =====================================================================
# HEADER
# =====================================================================
h1("CypherPro — Stripe / NMI Migration Q&A")
sub = doc.add_paragraph()
r = sub.add_run("Response to your discovery questions. February 2026.")
r.italic = True
r.font.color.rgb = SLATE
r.font.size = Pt(10.5)

spacer()


# =====================================================================
# CURRENT STRIPE SETUP
# =====================================================================
h2("CURRENT STRIPE SETUP")

q(1, "Is your current Stripe integration built on Stripe Connect "
     "(platform/sub-merchant model), or individual Stripe API keys per "
     "customer? This affects how much lift the switch to individual NMI "
     "merchant accounts will be on your end.")
a("We run a single-account platform model — one Stripe account, one API "
  "key, and every customer is a Stripe Customer object inside that account. "
  "We are NOT on Stripe Connect (no sub-merchants, no application fees, no "
  "transfer_data) and we do not hold per-customer API keys. Affiliate "
  "revenue-share is tracked as a ledger entry in our own database and "
  "settled off-Stripe, since we do not need Connect for that today. "
  "Practical implication: moving to per-merchant NMI accounts is a real "
  "migration for us, but not because we have Connect logic to unwind — "
  "rather because we currently have zero per-customer merchant-account "
  "abstraction at all. We will need to add one either way.")

q(2, "Beyond the customer-facing bill-pay / Payments AI feature, does "
     "Stripe also power your own CypherPro subscription billing today "
     "(the $38–$149/mo plans)?")
a("Yes — Stripe powers 100% of our own subscription billing. Worth "
  "flagging one thing that came up in review: the customer-facing "
  "bill-pay / Payments AI feature is on the roadmap but NOT yet in "
  "production. Currently the only Stripe usage anywhere in the app is "
  "our own subscription billing for the $38 / $79 / $95 / $149 / $349 "
  "plans. When we build customer bill-pay we will need to decide together "
  "whether it flows through NMI or Stripe, and that decision should be "
  "part of the scoping conversation.")


# =====================================================================
# REPLACING YOUR OWN SAAS BILLING
# =====================================================================
h2("REPLACING YOUR OWN SAAS BILLING")

q(3, "What are you using for subscription billing today — Stripe's "
     "Billing / Subscriptions API (proration, plan changes, and dunning "
     "built in), or Stripe Payment Links per plan? This tells us how much "
     "subscription-management logic needs to be rebuilt vs. carried over "
     "as-is.")
a("Full Stripe Billing — Subscriptions API + Invoices API, not Payment "
  "Links. Specifically:")
bullets([
    "New signups flow through stripe.checkout.Session.create(mode="
    "\"subscription\", …) so we can attach affiliate attribution and "
    "company metadata (Payment Links do not support that).",
    "Enterprise / flat-rate accounts are billed monthly via "
    "stripe.Invoice.create + InvoiceItem + finalize_invoice on a scheduler.",
    "Proration, plan changes, and dunning are all handled natively by "
    "Stripe — we did not reimplement any of it.",
    "State sync is webhook-driven and idempotent — "
    "checkout.session.completed, invoice.paid, "
    "customer.subscription.updated / deleted.",
    "Prices live in Stripe and are referenced by env-var Price IDs, so we "
    "can change pricing without a code deploy.",
])
a("Migration impact: swapping to a different processor means rebuilding "
  "the subscription-management layer we currently inherit from Stripe "
  "(proration math, dunning cadence, webhook event surface). Non-trivial, "
  "but scoped.")

q(4, "How does the 14-day free trial work today, specifically the "
     "\"no credit card required\" part — is a card collected at signup "
     "and just not charged, or not collected at all until day 14? We "
     "would need to replicate whichever flow that is.")
a("Card IS collected at signup and just not charged during the 14-day "
  "trial. On day 15, unless the customer has canceled, we auto-charge "
  "and convert them to paid. Our marketing copy currently uses the "
  "phrase \"no card required until day 15,\" which means \"you will not "
  "be billed for 14 days\" — not \"we do not ask for a card at signup.\" "
  "That language is worth cleaning up on our end for clarity, but the "
  "operational flow you need to replicate is: collect card at signup, "
  "hold, auto-bill on day 15 unless canceled.")

q(5, "Does your app need to know subscription status in real time, e.g. "
     "to gate feature access by plan? If so, which events matter most — "
     "successful charge, failed charge, cancellation, plan change?")
a("Yes — subscription status drives real-time feature access, and it is "
  "wired end-to-end via Stripe webhooks. We keep a per-company "
  "billing_state field (pending / active / past_due / canceled) as the "
  "source of truth for a blocking modal on the frontend, plus per-user "
  "subscription_status for the \"My Billing\" surface. Both flip on "
  "webhook, not on polling.")
a("Ranked by criticality for our product:")
bullets([
    "checkout.session.completed — unlock a new customer (highest priority).",
    "invoice.payment_failed — throw up the blocking modal, protect revenue.",
    "invoice.paid — restore access after past-due, credit the affiliate "
    "on the referral ledger.",
    "customer.subscription.updated / deleted — plan changes and clean "
    "cancellations.",
])
a("Every event is written to a stripe_webhook_events collection before we "
  "act on it, so aggressive retries cannot create duplicates. Failed "
  "handlers still return 200 so retries stop, and the raw event stays "
  "available for manual replay.")

q(6, "On pricing for your own subscription charges specifically, two "
     "paths worth deciding between: (a) Dual pricing — each plan tier "
     "shows two prices (a card price and a slightly lower ACH/bank-"
     "transfer price), same model as merchant dual pricing, with your "
     "subscriber paying the processing cost directly; or (b) CypherPro "
     "absorbs the fee — your advertised price per tier stays exactly as "
     "it is today, and we charge CypherPro a standard processing rate "
     "instead, which would come in below what you are paying Stripe today "
     "since we are removing their markup from the stack. (a) keeps your "
     "margin exactly where it is today but adds a second price to your "
     "pricing page; (b) keeps your pricing page unchanged but means "
     "eating a processing cost, albeit a lower one than you are paying "
     "now. Let me know which direction you guys like the most!")
a("Still finalizing this decision. I will commit to a direction within "
  "two weeks — the trade-off between pricing-page complexity and margin "
  "needs a bit more thought and a check with a couple of prospective "
  "Practice Partners before we lock in the shape of the pricing page. "
  "Will follow up directly with a decision.")

q(7, "(Question 7 was missing from the numbered list in the document I "
     "received — please confirm whether Q7 was intentionally skipped or "
     "if I missed a page.)")
a("Awaiting Q7 from you. Happy to answer it as soon as it comes through.")


# =====================================================================
# VOLUME & MERCHANT DATA
# =====================================================================
h2("VOLUME & MERCHANT DATA")

q(8, "Is the 100–500 merchant / $25MM-month figure current live Stripe "
     "processing volume, or a rollout target? From our first call I took "
     "it as projected once CypherPro goes live. If that is the case — "
     "what is the current monthly figure line? We want to size "
     "underwriting and staffing correctly either way, but need to know "
     "which one we are planning against.")
a("The $25MM / month is a rollout target, not current live volume. The "
  "math behind it: at 5,000 software users, each invoicing their own "
  "clients an average of $5,000 / month, we would land at ~$25MM / month "
  "processed through the platform.")
a("Current live processing volume through the software for customer "
  "bill-pay is $0 — that feature is not in production yet (see Q2). "
  "Our current Stripe volume is only our own SaaS subscription "
  "revenue, which is a fundamentally different traffic profile from "
  "what we are sizing this migration against.")

q(9, "Could you share anonymized/aggregate data on your current base — "
     "transaction count, average processing volume per customer, and "
     "card vs. ACH mix? This directly shapes the fee structure and lets "
     "us model this accurately together rather than guessing.")
a("Because customer bill-pay is not yet live, we do not have "
  "representative merchant transaction data to share on the payments "
  "side. Happy to share anonymized data on our own subscription-billing "
  "volume if that is useful for any reason, but it will not "
  "meaningfully inform the fee model for merchant payments. Once we "
  "have three months of live merchant activity post-launch, we can "
  "revisit the fee structure with real data.")

q(10, "Any sense of which business types/industries make up most of "
      "your customer base? Some categories need extra underwriting "
      "review, and knowing that now avoids surprises during the "
      "migration window.")
a("Across a broad set of legitimate small-business verticals — "
  "restaurants, professional services (CPAs, bookkeepers, consultants), "
  "agencies, contractors, retail, ecommerce, and healthcare admin. We "
  "are NOT planning to serve traditionally high-risk verticals: no "
  "adult, no gambling, no CBD/cannabis, no timeshare, no MLM, no "
  "firearms, no crypto exchanges. That should keep underwriting review "
  "focused on standard-risk SIC codes.")


# =====================================================================
# AUTOMATED APPLICATION — RISK-BASED DOCUMENT TRIGGERS
# =====================================================================
h2("AUTOMATED APPLICATION — RISK-BASED DOCUMENT TRIGGERS")

q(11, "Is [risk-based document trigger logic in the application form] "
      "something your application intake logic could capture and act on? "
      "This isn't a required — but it would help reduce a lot of back "
      "and forth and reduce the overall time from application to live.")
a("Absolutely — this is a natural fit for our application intake flow "
  "and we can automate it out of the gate. Conditional logic that "
  "auto-requests 3 months of business bank statements or two paid "
  "customer invoices when a merchant flags themselves as CNP / B2B is "
  "straightforward. The \"prior-processing statements\" trigger is "
  "worth building too, though we will want to soft-fail gracefully — "
  "many small-business owners genuinely do not know how to pull those "
  "statements, so we will offer an inline how-to and let underwriting "
  "manually request only if truly needed. This capability will keep "
  "improving as we see which triggers reduce actual back-and-forth "
  "the most.")


# =====================================================================
# INTEGRATION SCOPE
# =====================================================================
h2("INTEGRATION SCOPE")

q(12, "On OptiPay360 — when you mentioned possibly integrating it "
      "instead of building your own, could you say more about what you "
      "are picturing? I think this could be great but want to get a "
      "better idea of what you are thinking and to be sure we have the "
      "right support team on top of it and ready for the volume. The "
      "last thing we will do is reduce the value of the support we "
      "offer to you guys or any customers on the platform.")
a("Let me come back on this one with a clearer picture. I want to talk "
  "through the OptiPay360 architecture internally before locking in the "
  "shape of the integration — specifically what we would ship natively "
  "in CypherPro vs. what we would rely on OptiPay360 for. Once I have a "
  "clean scope written up, I will send it over so we can align on "
  "support ownership before anything gets built.")


# =====================================================================
# PAYMENTS UX & FEES (MERCHANT SIDE)
# =====================================================================
h2("PAYMENTS UX & FEES (MERCHANT SIDE)")

q(13, "For ACH specifically: would your team be open to a disclosed "
      "pass-through fee shown to the payer at time of payment (similar "
      "to how card dual pricing already works), or does \"no hard "
      "monthly fees\" mean no visible payment fees anywhere in the flow?")
a("Yes — we are open to a disclosed ACH pass-through fee shown to the "
  "payer at time of payment, same pattern as card dual pricing.")
a("Clarifying what we mean by \"no hard monthly fees\": we do not want "
  "the merchant to be charged a monthly fee in a month when they process "
  "nothing. The pattern we are trying to avoid is: merchant gets "
  "approved, has an idle month with zero transactions, and still gets "
  "hit with a mandatory $35 monthly service fee on top of their CypherPro "
  "subscription. Fees at time of transaction (whether merchant-absorbed "
  "or payer-disclosed) are fine. Fixed monthly access fees regardless of "
  "usage are what we want to avoid on the payments layer.")


# =====================================================================
# SECURITY
# =====================================================================
h2("SECURITY")

q(14, "Given your SOC 2 Type II posture, do you want card data kept "
      "fully out of your own systems (hosted fields, similar to "
      "Collect.js), or are you planning to collect fields within your "
      "own UI?")
a("We want card data kept fully out of our own systems — hosted fields "
  "are exactly the model we prefer. This is a hard requirement, driven "
  "by our SOC 2 posture and by our decision to stay in PCI SAQ-A scope "
  "rather than SAQ-D. Practically: post-migration, our plan is to route "
  "cardholder data straight to NMI via hosted fields (Collect.js pattern) "
  "so nothing sensitive ever touches our infrastructure. Just to set "
  "expectations on the current state — the NMI hosted-fields integration "
  "is not built yet in the app; the code we ship today only touches "
  "Stripe. Wiring the hosted-fields flow is part of the migration work.")


# =====================================================================
# WRAP
# =====================================================================
h2("SUMMARY OF OPEN ITEMS")
a("Three things I owe you a follow-up on, plus one clarifying question "
  "on my end:")
bullets([
    "Q6 — pricing direction (dual pricing vs. CypherPro absorbs the fee): "
    "coming within 2 weeks.",
    "Q7 — please resend, this question appears to be missing from the "
    "numbered list I received.",
    "Q12 — OptiPay360 integration scope: coming in a separate write-up "
    "once I've talked it through internally.",
    "Q9 — will share anonymized merchant data 3 months after customer "
    "bill-pay goes live.",
])

# =====================================================================
# Build
# =====================================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
